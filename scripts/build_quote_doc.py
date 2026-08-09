from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from build_word_docs import (
    CODE_FILL,
    GRID,
    INK,
    LIME,
    MUTED,
    PAPER,
    TABLE_FILL,
    WHITE,
    add_page_field,
    add_rich_text,
    configure_document,
    set_cell_margins,
    set_cell_shading,
    set_font,
    set_table_borders,
    set_table_geometry,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "word" / "cotizacion-esacob-apl.docx"


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag.endswith("}pPr"):
            continue
        paragraph._p.remove(child)


def configure_quote(doc):
    configure_document(doc, compact=False)
    section = doc.sections[0]

    header = section.header
    hp = header.paragraphs[0]
    clear_paragraph(hp)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("ESACOB SERVICIOS GENERALES S.A.C.  |  Cotización comercial")
    set_font(r, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    clear_paragraph(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("COT-ESACOB-2026-008  |  Página ")
    set_font(r, size=8.5, color=MUTED)
    add_page_field(fp)
    return section


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    set_font(r, size=10, color=MUTED, bold=True)
    return p


def add_title(doc):
    add_kicker(doc, "Propuesta comercial")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("Cotización comercial")
    set_font(r, size=28, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    r = p.add_run("Solución integral de seguimiento y gestión de flota")
    set_font(r, size=13.5, color=MUTED)


def add_metadata(doc):
    rows = [
        ("Cotización", "COT-ESACOB-2026-008"),
        ("Fecha de emisión", "04 de agosto de 2026"),
        ("Vigencia", "30 días"),
        ("Proveedor", "ESACOB SERVICIOS GENERALES S.A.C. | RUC 20610791132"),
        ("Domicilio", "Calle Poe, Mz. K, Lote 7, Surquillo, Lima"),
        ("Cliente", "APL Servicios Generales E.I.R.L. | RUC 20600397789"),
        ("Domicilio del cliente", "Lima / Lima / Ate, Perú"),
        ("Contacto", "manuel_cortezballardo@outlook.com | 986 969 037"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [2200, 7160])
    set_table_borders(table, color=GRID, size="4")
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], TABLE_FILL)
        for cell in row.cells:
            set_cell_margins(cell, top=80, start=120, bottom=80, end=120)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_font(r, size=9.5, color=INK, bold=True)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_font(r, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_total_band(doc):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [6000, 3360])
    set_table_borders(table, color=LIME, size="8")
    left, right = table.rows[0].cells
    set_cell_shading(left, INK)
    set_cell_shading(right, LIME)
    for cell in (left, right):
        set_cell_margins(cell, top=140, start=160, bottom=140, end=160)
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("TOTAL DE LA PROPUESTA")
    set_font(r, size=10, color=LIME, bold=True)
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("USD 20,000 + IGV")
    set_font(r, size=19, color=WHITE, bold=True)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Total con IGV")
    set_font(r, size=9.5, color=INK, bold=True)
    p = right.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("USD 23,600")
    set_font(r, size=18, color=INK, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    add_rich_text(p, text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    add_rich_text(p, text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        add_rich_text(p, item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table, color=GRID, size="6")
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, TABLE_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(header)
        set_font(r, size=9.2, color=INK, bold=True)
    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_rich_text(p, value, base_size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_note(doc, label, text, fill=TABLE_FILL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=GRID, size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=100, start=140, bottom=100, end=140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_font(r, size=9.8, color=INK, bold=True)
    r = p.add_run(text)
    set_font(r, size=9.8, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def build():
    doc = Document()
    configure_quote(doc)
    add_title(doc)
    add_metadata(doc)
    add_total_band(doc)
    add_note(
        doc,
        "Resumen ejecutivo",
        "ESACOB implementará y entregará una solución integral de seguimiento y gestión de flota para APL, con una capacidad de referencia de 100 vehículos y aproximadamente 2,000 viajes mensuales.",
    )

    add_heading(doc, "1. Objeto de la propuesta")
    add_body(
        doc,
        "Implementación, personalización, publicación y transferencia de una solución de gestión de flota compuesta por aplicaciones Android para pasajeros y conductores, dashboard web de operación y backend cloud sobre Firebase y Google Cloud.",
    )

    add_heading(doc, "2. Alcance incluido")
    add_table(
        doc,
        ["Componente", "Incluye"],
        [
            ("App de pasajeros Android", "Registro, solicitud de viajes inmediatos y programados, visualización del vehículo asignado, seguimiento del viaje e historial."),
            ("App de conductores Android", "Registro, carga de documentos, aprobación, turno, envío de ubicación GPS, recepción de viajes y ejecución por estados."),
            ("Dashboard web", "Mapa de flota, conductores, lugares, viajes, asignaciones, sedes, usuarios, historial y monitoreo operativo."),
            ("Backend y servicios cloud", "Firebase Realtime Database, Authentication, Storage, Cloud Functions y Cloud Messaging; configuración, despliegue y pruebas."),
            ("Mapas y rutas", "Mapbox Maps SDK/GL JS, Geocoding, Directions, Matrix y Traffic; tokens restringidos y alertas de consumo."),
            ("Personalización de marca", "Implementación del logo, colores y elementos de marca entregados por APL."),
            ("Publicación Android", "Preparación y publicación de las aplicaciones en Google Play Store, usando las cuentas y accesos de APL."),
            ("Documentación y capacitación", "Dossier comercial, manual de operación, referencia técnica, guía de despliegue y sesiones de capacitación."),
            ("Código fuente y transferencia", "Entrega del repositorio, documentación técnica y archivos necesarios después del pago final."),
            ("Administración inicial", "Administración, mantenimiento y soporte operativo por parte de ESACOB hasta el 31 de diciembre de 2026."),
        ],
        [2450, 6910],
    )

    add_heading(doc, "3. Capacidad e intervalos de referencia")
    add_table(
        doc,
        ["Referencia", "Valor"],
        [
            ("Flota inicial", "100 vehículos"),
            ("Volumen estimado", "Aproximadamente 2,000 viajes mensuales"),
            ("Envío de ubicación del conductor", "Inmediato al iniciar y luego cada 15 segundos"),
            ("Actualización del dashboard", "Cambios en tiempo real y revisión visual cada 10 segundos"),
            ("Ubicación considerada obsoleta", "Después de 45 segundos sin actualización"),
            ("Retiro visual por desconexión prolongada", "Después de 3 minutos"),
            ("Búsqueda de viajes del pasajero", "Consulta aproximada cada 3 segundos"),
            ("Consulta del viaje activo", "Estado cada 4 segundos y posición del conductor cada 5 segundos"),
        ],
        [3300, 6060],
    )

    add_heading(doc, "4. Propuesta económica")
    add_table(
        doc,
        ["Concepto", "Importe"],
        [
            ("Implementación, personalización, publicación, documentación, transferencia del código fuente y administración inicial hasta el 31/12/2026", "USD 20,000.00"),
            ("IGV 18%", "USD 3,600.00"),
            ("Total de la propuesta", "USD 23,600.00"),
        ],
        [6900, 2460],
    )

    add_heading(doc, "5. Forma de pago")
    add_table(
        doc,
        ["Hito", "Base", "IGV", "Total"],
        [
            ("Al firmar la propuesta", "USD 4,000.00", "USD 720.00", "USD 4,720.00"),
            ("Pago final, a más tardar el 30/09/2026", "USD 16,000.00", "USD 2,880.00", "USD 18,880.00"),
            ("Total", "USD 20,000.00", "USD 3,600.00", "USD 23,600.00"),
        ],
        [3550, 1800, 1700, 2310],
    )
    add_note(
        doc,
        "Condición de transferencia",
        "La propiedad definitiva del código fuente, repositorio y archivos de entrega se transferirá a APL una vez recibido el pago final.",
        fill=PAPER,
    )

    add_heading(doc, "6. Infraestructura, servidores y APIs")
    add_body(
        doc,
        "Los servicios cloud y de terceros serán contratados y pagados directamente por APL. ESACOB configurará y pondrá en marcha el entorno utilizando cuentas propiedad de APL, para que el cliente conserve el control de sus datos, facturación y accesos.",
    )
    add_bullets(
        doc,
        [
            "APL asumirá directamente los consumos de Firebase, Google Cloud, Mapbox, hosting, almacenamiento, dominios y otros servicios de terceros.",
            "ESACOB realizará la configuración de proyectos, claves restringidas, permisos, alertas de presupuesto, despliegue y validación técnica.",
            "Los celulares o dispositivos GPS, líneas de datos y accesorios ya son responsabilidad de APL y no forman parte de esta cotización.",
            "Las cuentas de Google Play Store y sus cargos de publicación serán gestionados por APL.",
            "El consumo de APIs no es ilimitado ni está incluido como gasto operativo de ESACOB dentro del precio de implementación.",
        ],
    )

    add_heading(doc, "7. Cronograma y aceptación")
    add_bullets(
        doc,
        [
            "Inicio: después de la firma y del pago inicial del 20%.",
            "APL entregará logo, colores, accesos, información de operación y permisos necesarios para la publicación.",
            "Entrega operativa máxima: 13 de agosto de 2026.",
            "La solución será preparada para la operación de APL durante el evento de septiembre de 2026.",
            "La conformidad será emitida por Alexandra Gisell Arias Espinoza, gerente general de APL.",
            "El pago final tendrá como fecha máxima el 30 de septiembre de 2026.",
        ],
    )

    add_heading(doc, "8. Soporte, administración y límites")
    add_body(
        doc,
        "La administración, mantenimiento y soporte operativo de ESACOB están incluidos hasta el 31 de diciembre de 2026. Este alcance cubre corrección de errores, mantenimiento técnico, operación y asistencia sobre las funciones incluidas.")
    add_bullets(
        doc,
        [
            "Los nuevos módulos, funciones no descritas, integraciones adicionales o cambios de alcance se cotizarán por separado.",
            "La continuidad del soporte después del 31 de diciembre de 2026 será opcional y se definirá mediante un nuevo acuerdo.",
            "La fecha de entrega depende de que APL entregue oportunamente accesos, contenido de marca, cuentas de publicación y datos de operación.",
        ],
    )

    add_heading(doc, "9. Aprobación")
    add_table(
        doc,
        ["Por ESACOB", "Por APL"],
        [
            ("Gino Carlos Alexander Cortez Ballardo\nGerente General\nESACOB SERVICIOS GENERALES S.A.C.", "Alexandra Gisell Arias Espinoza\nGerente General\nAPL Servicios Generales E.I.R.L."),
            ("Firma: ______________________________\nFecha: ______________________________", "Firma: ______________________________\nFecha: ______________________________"),
        ],
        [4680, 4680],
    )
    add_note(
        doc,
        "Observación comercial",
        "Esta cotización tiene una vigencia de 30 días desde el 04 de agosto de 2026 y deberá complementarse con un contrato de prestación de servicios y transferencia de propiedad intelectual.",
        fill=CODE_FILL,
    )

    doc.core_properties.title = "Cotización ESACOB - APL"
    doc.core_properties.subject = "Propuesta comercial de solución de seguimiento y gestión de flota"
    doc.core_properties.author = "ESACOB SERVICIOS GENERALES S.A.C."
    doc.core_properties.keywords = "cotización, fleet tracking, APL, ESACOB"
    doc.core_properties.comments = "Prepared for commercial review."
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    build()
