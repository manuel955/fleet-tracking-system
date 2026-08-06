from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "docs"
OUTPUT_DIR = SOURCE_DIR / "word"

INK = "081618"
LIME = "C8F267"
PAPER = "F5F7F3"
MUTED = "5B686A"
GRID = "D7E1D8"
TABLE_FILL = "E8F3C8"
CODE_FILL = "EEF3F0"
WHITE = "FFFFFF"

DOC_METADATA = {
    "README.md": ("Índice de documentación", "Mapa de lectura, alcance y pendientes de entrega"),
    "dossier-comercial.md": ("Dossier comercial", "Documento para presentar la solución al comprador"),
    "manual-operacion.md": ("Manual de operación", "Guía para pasajeros, conductores y operadores"),
    "referencia-tecnica.md": ("Referencia técnica", "Arquitectura, datos, estados e integraciones"),
    "guia-despliegue-y-entrega.md": ("Guía de despliegue y entrega", "Instalación, publicación, validación y transferencia"),
}


def set_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name="Calibri", size=11, color=INK, bold=None):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            tr_pr = row._tr.get_or_add_trPr()
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=GRID, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def configure_document(doc, compact=True):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6 if compact else 8)
    normal.paragraph_format.line_spacing = 1.25 if compact else 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, INK, 18 if compact else 18, 10 if compact else 10),
        ("Heading 2", 13, INK, 14 if compact else 12, 7 if compact else 6),
        ("Heading 3", 12, MUTED, 10 if compact else 8, 5 if compact else 4),
    ):
        style = styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        set_style_font(style, size=11, color=INK)
        style.paragraph_format.left_indent = Inches(0.5 if compact else 0.375)
        style.paragraph_format.first_line_indent = Inches(-0.25 if compact else -0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25 if compact else 1.167

    code = styles.add_style("Code Block", 1)
    set_style_font(code, name="Consolas", size=9, color=INK)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.15)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.0

    callout = styles.add_style("Callout", 1)
    set_style_font(callout, size=10.5, color=INK)
    callout.paragraph_format.left_indent = Inches(0.15)
    callout.paragraph_format.right_indent = Inches(0.15)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.15

    # Quiet running header and footer.
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("APL Logistic  |  Documentación de la solución")
    set_font(hr, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fr = fp.add_run("Documento de entrega  |  Página ")
    set_font(fr, size=8.5, color=MUTED)
    add_page_field(fp)

    return section


def clean_inline(text):
    text = re.sub(r"!\[([^]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r"\[([^]]+)\]\([^)]*\)", r"\1", text)
    text = text.replace("`", "")
    return text


def add_rich_text(paragraph, text, base_size=11):
    # Preserve the useful inline emphasis from Markdown without exposing Markdown syntax.
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(clean_inline(text[cursor:match.start()]))
            set_font(run, size=base_size, color=INK)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(clean_inline(token[2:-2]))
            set_font(run, size=base_size, color=INK, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Consolas", size=base_size - 1, color=INK)
        else:
            run = paragraph.add_run(clean_inline(token[1:-1]))
            set_font(run, size=base_size, color=INK, italic=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(clean_inline(text[cursor:]))
        set_font(run, size=base_size, color=INK)


def add_title_block(doc, title, subtitle, source_name):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("APL LOGISTIC")
    set_font(r, size=10, color=MUTED, bold=True)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(6)
    title_p.paragraph_format.keep_with_next = True
    r = title_p.add_run(title)
    set_font(r, size=28, color=INK, bold=True)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(14)
    r = sub_p.add_run(subtitle)
    set_font(r, size=13, color=MUTED)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(14)
    r = meta.add_run("Preparado para entrega comercial y técnica  |  Fuente: ")
    set_font(r, size=9.5, color=MUTED, bold=True)
    r = meta.add_run(f"{source_name}  |  Auditoría estática: 03/08/2026")
    set_font(r, size=9.5, color=MUTED)

    band = doc.add_table(rows=1, cols=1)
    set_table_geometry(band, [9360])
    set_table_borders(band, color=GRID, size="4")
    cell = band.cell(0, 0)
    set_cell_shading(cell, PAPER)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run("Nota: los valores descritos reflejan el comportamiento implementado. La latencia final también depende del GPS, la cobertura y los servicios externos.")
    set_font(run, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table(lines):
    rows = []
    for line in lines:
        stripped = line.strip().strip("|")
        cells = [clean_inline(cell.strip()) for cell in stripped.split("|")]
        rows.append(cells)
    if len(rows) > 1 and all(re.fullmatch(r"\s*:?-{3,}:?\s*", cell) for cell in rows[1]):
        rows.pop(1)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    normalized = [row + [""] * (cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=cols)
    if cols == 1:
        widths = [9360]
    elif cols == 2:
        widths = [2700, 6660]
    elif cols == 3:
        widths = [1900, 3500, 3960]
    elif cols == 4:
        widths = [1450, 2450, 2550, 2910]
    else:
        base = 9360 // cols
        widths = [base] * cols
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    set_table_borders(table)
    for row_index, row in enumerate(normalized):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            if row_index == 0:
                set_cell_shading(cell, TABLE_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_rich_text(p, value, base_size=9 if len(normalized) > 8 else 9.5)
            if row_index == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph(style="Code Block")
        p.paragraph_format.keep_together = True
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), CODE_FILL)
        pPr.append(shd)
        r = p.add_run(line)
        set_font(r, name="Consolas", size=8.5, color=INK)


def add_body_from_markdown(doc, text):
    lines = text.splitlines()
    first_h1 = True
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# ") and first_h1:
            first_h1 = False
            i += 1
            continue
        first_h1 = False

        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, parse_table(table_lines))
            continue

        match = re.match(r"^(#{2,3})\s+(.+)$", stripped)
        if match:
            level = len(match.group(1)) - 1
            p = doc.add_paragraph(style=f"Heading {level}")
            add_rich_text(p, match.group(2))
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Callout")
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), TABLE_FILL)
            pPr.append(shd)
            add_rich_text(p, stripped[1:].strip(), base_size=10.5)
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            item = re.sub(r"^[-*]\s+", "", stripped)
            add_rich_text(p, item)
            i += 1
            continue

        if re.match(r"^\d+[.)]\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            item = re.sub(r"^\d+[.)]\s+", "", stripped)
            add_rich_text(p, item)
            i += 1
            continue

        if stripped in {"---", "***"}:
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith(("# ", "## ", "### ", ">", "|", "```")):
                break
            if re.match(r"^[-*]\s+", nxt) or re.match(r"^\d+[.)]\s+", nxt) or nxt in {"---", "***"}:
                break
            paragraph_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_rich_text(p, " ".join(paragraph_lines))


def build_doc(source_name):
    source_path = SOURCE_DIR / source_name
    raw = source_path.read_text(encoding="utf-8")
    title, subtitle = DOC_METADATA[source_name]
    doc = Document()
    configure_document(doc, compact=source_name != "dossier-comercial.md")
    add_title_block(doc, title, subtitle, source_name)
    add_body_from_markdown(doc, raw)

    doc.core_properties.title = title
    doc.core_properties.subject = subtitle
    doc.core_properties.author = "APL Logistic"
    doc.core_properties.keywords = "fleet tracking, fleet operations, documentation"
    doc.core_properties.comments = "Prepared from the audited project documentation."

    output_name = source_path.stem + ".docx"
    output_path = OUTPUT_DIR / output_name
    doc.save(output_path)
    return output_path


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for source_name in DOC_METADATA:
        build_doc(source_name)
    print(f"Created {len(DOC_METADATA)} DOCX files in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
